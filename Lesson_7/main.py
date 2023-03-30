# 1. Создание текстового файла с данными
with open("car_data.txt", "w") as f:
    f.write("Марка авто,Модель авто,Расход топлива,Стоимость\n")
    f.write("Toyota,Camry,7.8,1500000\n")
    f.write("Kia,Sportage,8.5,1200000\n")
    f.write("Hyundai,Accent,6.2,900000\n")

# 2. Создание шаблона документа doc
from docx import Document

document = Document()
document.add_heading("Данные автомобилей")
document.add_paragraph("Марка авто: ...")
document.add_paragraph("Модель авто: ...")
document.add_paragraph("Расход топлива: ...")
document.add_paragraph("Стоимость: ...")

document.save("car_data.docx")

# 3. Внесение данных из файла в шаблон
import csv

with open("car_data.csv", "w") as f:
    writer = csv.writer(f)
    writer.writerow(["Марка авто", "Модель авто", "Расход топлива", "Стоимость"])
    with open("car_data.txt", "r") as data:
        for line in data:
            row = line.strip().split(",")
            writer.writerow(row)

# 4. Создание json-файла с данными о машине
import json

data = []
with open("car_data.txt", "r") as f:
    header = f.readline().strip().split(",")
    for line in f:
        row = line.strip().split(",")
        car_data = {header[i]:row[i] for i in range(len(header))}
        data.append(car_data)

with open("car_data.json", "w") as f:
    json.dump(data, f)